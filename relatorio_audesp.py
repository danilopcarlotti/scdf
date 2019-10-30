import pymongo, pandas as pd, pickle, os, re
from docx import Document

myclient = pymongo.MongoClient('mongodb://localhost:27017/')
mydb = myclient['Audesp']
entidades = mydb['entidades']
licitacoes = mydb['licitacoes']
licitantes = mydb['licitantes']
municipios = mydb['municipios']

db_obitos = myclient['Pessoas']
pessoas_mortas = db_obitos['pessoas_fisicas']

db_rf = myclient["RF"]
empresas = db_rf["Empresas"]

path_relatorios = ''

# MATRIZ DE RISCO

def find_variaveis_receita_federal(cnpj_licitante):
    dicionario = empresas.find_one({"_id":cnpj_licitante})
    if dicionario:
        return (dicionario['situacao_cadastral'],dicionario['data_inicio_atividade'],dicionario['porte_empresa'])
    else:
        return (-1,-1,-1)

# def empresa_loser(cnpj_licitante):
#     return False

# def ganha_loser(cnpj_licitante):
#     return (False, [])

# def sem_competicao(id_empresa):
#     return (False, [])

# def valor_acumulado_empresa(id_empresa):
#     return 0.0

def pessoa_morta(cpf):
	cpf = cpf.replace('/','').replace('-','').replace('.','')
	dados_morto = pessoas_mortas.find_one({'_id':str(cpf)})
	if dados_morto:
		return dados_morto
	return False

def socio_servidor_publico(cpf):
	return False

#  MATRIZ DE RISCO

def consertar_str_cnpj(cnpj):
    cnpj = cnpj.replace('/','').replace('-','').replace('.','')
    if len(cnpj) == 14:
        return cnpj
    elif len(cnpj) > 14:
        return consertar_str_cnpj(cnpj[:-1])
    else:
        return consertar_str_cnpj('0'+cnpj)

def dados_entidade(id_entidade):
	dic_ent = entidades.find_one({'_id':id_entidade})
	if dic_ent:
		dic_aux = {'nome_completo':dic_ent['nome_completo']}
		dados_municipio = dados_municipio(dic_ent['municipio_id'])
		return [dic_aux,dados_municipio]
	return False

def dados_licitantes(cnpj_licitante):
	dic_licitantes = licitantes.find_one({'_id':cnpj_licitante})
	if dic_licitantes:
		dic_aux = {}
		colunas_interesse = ['cpf_adm','nome_licitante','licitacao_id','socios']
		for c in colunas_interesse:
			dic_aux[c] = dic_licitantes[c]
		return dic_aux
	return False

def dados_municipio(municipio_id):
	municipio = municipios.find_one({'municipio_id':municipio_id})
	if municipio:
		return municipio
	return False

def dados_licitacao(id_licitacao):
	dic_licitacao = licitacoes.find_one({'_id':id_licitacao})
	if dic_licitacao:
		dic_aux = {}
		colunas_interesse = ['ano_licitacao','descricao_objeto','valor_total']
		for c in colunas_interesse:
			dic_aux[c] = dic_licitacao[c]
		entidade, dados_mun = dados_entidade(dic_licitacao['entidade_id'])
		return [id_licitacao,dic_aux,entidade,dados_mun]
	return (False,False,False,False)

def relatorio_cnpj(cnpj):
	cnpj_licitante = consertar_str_cnpj(cnpj)
	dados_licitante = dados_licitantes(cnpj_licitante)
	if not dados_licitante:
		document = Document()
		document.add_paragraph('Não foi encontrada nenhuma informação para a empresa')
		document.save(path_relatorios+'relatório_%s.docx' % (cnpj_licitante,))
		return
	dados_licitacoes = []
	for l_id in dados_licitante['licitacao_id']:
		dados_licitacoes.append(dados_licitacao(str(l_id)))
	texto = 'A empresa com CNPJ %s e razão social %s possui as seguintes informações, segundo consulta à base da RF que precisam ser validadas em tempo real:\n\n' % (str(cnpj_licitante),dados_licitante['nome_licitante'])
	dados_rf = find_variaveis_receita_federal(cnpj_licitante)
	if dados_rf:
		texto += '\tSituação cadastral: %s\n' % (dados_rf['situacao_cadastral'],)
		texto += '\tInício da atividade da empresa: %s\n' % (dados_rf['data_inicio_atividade'],)
		texto += '\tPorte da empresa: %s\n' % (dados_rf['porte_empresa'],)
	else:
		texto += '\tNão foram encontradas informações na base da RF para a empresa.\n'
	socios = dados_licitante['socios']
	if len(socios):
		socios = list(set(socios))
		texto += '\n\tA empresa possui os seguintes sócios:\n'
		for socio in socios:
			texto += '\tNúmero do documento do sócio: %s\n' % (str(socio),)
	if len(dados_licitacoes):
		texto += '\tRELATÓRIO GERAL\n\n\tA empresa participou de %s licitações.\n\n' % (str(len(dados_licitacoes),))
		texto += '\tINFORMAÇÕES GERAIS SOBRE A EMPRESA:\n\n'
		if empresa_loser(cnpj_licitante):
			texto += '\tA empresa em questão perdeu todas as licitações das quais participou\n'
		ganha_l, empresas = ganha_loser(cnpj_licitante)
		if ganha_l:
			texto += '\tA empresa em questão participou de licitações com empresas que só perdem. São estas as empresas:\n'
			for emp in empresas:
				texto += '\t' + emp + '\n'
		ganha_s_c, lics = sem_competicao(cnpj_licitante)
		if ganha_s_c:
			texto += '\tA empresa em questão ganhou as seguintes licitações sem concorrência:\n'
			for l in lics:
				texto += '\t' + lics + '\n'
		valor_acum = valor_acumulado_empresa
		if valor_acum:
			texto += '\tA empresa tem ganhos de licitações no valor total de %s' % (str(valor_acum),)
		texto += '\tInformações sobre licitações em que a empresa concorreu:\n'
		for id_licitacao,dic_aux,entidade,dados_mun in dados_licitacoes:
			if id_licitacao:
				texto += '\n\n\tLicitação registrada no banco de dados Audesp com id: %s\n' % (str(id_licitacao))
				texto += '\tA licitação ocorreu no ano %s, tem como objeto "%s" e está registrada tendo valor total de R$%s\n' % (str(int(float(dic_aux['ano_licitacao']))),str(dic_aux['descricao_objeto']),str(dic_aux['valor_total']))
				if entidade:
					texto += '\tA entidade que realizou a licitação se denomina: %s\n' % (entidade['nome_completo'])
					if dados_mun:
						texto += '\tA entidade que realizou a licitação se localiza no município: %s\n' % (dados_mun['ds_municipio'])
	else:
		texto += '\tNão há informações sobre licitações em que a empresa concorreu.'
	document = Document()
	document.add_paragraph(texto)
	document.save(path_relatorios+'relatório_%s.docx' % (cnpj_licitante,))

def relatorio_socios(documento_socio):
    documento_socio = documento_socio.replace('.','').replace('-','')
    empresas_socio = []
    dics = licitantes.find({'cpf_adm':documento_socio})
    for dic in dics:
        empresas_socio.append((dic['_id'],dic['nome_licitante']))
    doc = Document() 
    doc.add_paragraph('A pessoa com documento %s é sócia administradora das seguintes empresas:\n\n' % (documento_socio,))
    if len(empresas_socio):
        for cnpj, e in empresas_socio:
            doc.add_paragraph(cnpj+' empresa de nome comercial '+e+'\n')
    else:
        doc.add_paragraph('Essa pessoa não é administradora de nenhuma empresa')
    dados_morto = pessoa_morta(documento_socio)
	if dados_morto:
		doc.add_paragraph('\n\tESSA PESSOA MORREU. Dados do óbito:\n')
		doc.add_paragraph('\nData do óbito:%s\n' % (dados_morto['Data_Obito']))
	doc.save(path_relatorios+'socio_empresas_cpf_%s.docx' % (documento_socio,))

def main(cnpj, cpf):
	cnpj = str(cnpj)
	cpf = str(cpf)
	relatorio_cnpj(cnpj)
	relatorio_socios(cpf)

if __name__ == '__main__':
	main('11311279000140','')