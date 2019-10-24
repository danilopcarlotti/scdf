# -*- coding: utf-8 -*-

# Form implementation generated from reading ui file 'banco_precos_ui.ui'
#
# Created by: PyQt5 UI code generator 5.13.0
#
# WARNING! All changes made in this file will be lost!


from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtWidgets import QFileDialog
from PyQt5.QtWidgets import QMessageBox

from docx import Document
import pandas as pd, pymongo, re, pickle, sys, numpy as np
from scipy.spatial import distance
from sklearn.feature_extraction.text import TfidfVectorizer
from remove_accents import remove_accents

import image_logo_tce, pickle

myclient = pymongo.MongoClient('mongodb://localhost:27017/')
mydb = myclient['BEC']
audespdb = myclient['Audesp']
banco_precos_col = audespdb['bancoPrecos']
path_relatorio = ''

class Ui_MainWindow(object):
    def setupUi(self, MainWindow):
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(847, 471)
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.textBrowser = QtWidgets.QTextBrowser(self.centralwidget)
        self.textBrowser.setGeometry(QtCore.QRect(10, 10, 541, 111))
        self.textBrowser.setObjectName("textBrowser")
        self.lineEdit = QtWidgets.QLineEdit(self.centralwidget)
        self.lineEdit.setGeometry(QtCore.QRect(20, 220, 531, 41))
        self.lineEdit.setText("")
        self.lineEdit.setObjectName("lineEdit")
        self.pushButton_2 = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton_2.setGeometry(QtCore.QRect(20, 270, 531, 91))
        self.pushButton_2.setObjectName("pushButton_2")
        self.label = QtWidgets.QLabel(self.centralwidget)
        self.label.setGeometry(QtCore.QRect(140, 160, 311, 51))
        self.label.setObjectName("label")
        self.textEdit = QtWidgets.QTextEdit(self.centralwidget)
        self.textEdit.setGeometry(QtCore.QRect(570, 10, 251, 111))
        self.textEdit.setObjectName("textEdit")
        MainWindow.setCentralWidget(self.centralwidget)
        self.statusbar = QtWidgets.QStatusBar(MainWindow)
        self.statusbar.setObjectName("statusbar")
        MainWindow.setStatusBar(self.statusbar)
        self.toolBar = QtWidgets.QToolBar(MainWindow)
        self.toolBar.setObjectName("toolBar")
        MainWindow.addToolBar(QtCore.Qt.TopToolBarArea, self.toolBar)
        self.actionSelecionar_pasta = QtWidgets.QAction(MainWindow)
        self.actionSelecionar_pasta.setObjectName("actionSelecionar_pasta")

        self.retranslateUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)

        self.pushButton_2.clicked.connect(self.gerar_relatorio)

    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "MainWindow"))
        self.textBrowser.setHtml(_translate("MainWindow", "<!DOCTYPE HTML PUBLIC \"-//W3C//DTD HTML 4.0//EN\" \"http://www.w3.org/TR/REC-html40/strict.dtd\">\n"
"<html><head><meta name=\"qrichtext\" content=\"1\" /><style type=\"text/css\">\n"
"p, li { white-space: pre-wrap; }\n"
"</style></head><body style=\" font-family:\'Garuda\'; font-size:11pt; font-weight:400; font-style:normal;\">\n"
"<p align=\"center\" style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><span style=\" font-size:12pt;\">Solução em ciência de dados forenses 1.0</span></p>\n"
"<p align=\"center\" style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><span style=\" font-size:12pt;\">Banco de preços - Bolsa Eletrônica de Compras (BEC)</span></p>\n"
"<p align=\"center\" style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><span style=\" font-size:12pt;\">Ministério Público do Estado de SP</span></p></body></html>"))
        self.pushButton_2.setText(_translate("MainWindow", "Gerar relatório dos 5 itens mais parecidos com o item pesquisado"))
        self.label.setText(_translate("MainWindow", "Escreva o nome do produto a ser pesquisado"))
        self.textEdit.setHtml(_translate("MainWindow", "<!DOCTYPE HTML PUBLIC \"-//W3C//DTD HTML 4.0//EN\" \"http://www.w3.org/TR/REC-html40/strict.dtd\">\n"
"<html><head><meta name=\"qrichtext\" content=\"1\" /><style type=\"text/css\">\n"
"p, li { white-space: pre-wrap; }\n"
"</style></head><body style=\" font-family:\'Garuda\'; font-size:11pt; font-weight:400; font-style:normal;\">\n"
"<p align=\"center\" style=\"-qt-paragraph-type:empty; margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><br /></p>\n"
"<p align=\"center\" style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"> <img src=\":/newPrefix/MPSPAtivo 5.png\" /></p></body></html>"))
        self.toolBar.setWindowTitle(_translate("MainWindow", "toolBar"))
        self.actionSelecionar_pasta.setText(_translate("MainWindow", "Selecionar pasta"))

    def give_recommendation(self,desc_item, threshold=0.75):
        words = []
        dicionario_banco_precos = pickle.load(open('dicionario_banco_precos.pickle','rb'))
        tfv_fit = pickle.load(open('tfv_fit.pickle','rb'))
        vec_word = tfv_fit.transform([desc_item]).A[0]
        for dic in banco_precos_col.find({}):
            if len(words) > 5:
                break
            else:
                vec_comp = np.array(dic['vetor_tfidf'])
                if distance.cdist([vec_word],[vec_comp],'cosine') < threshold:
                    words.append(dic['_id'])
        resultado = []
        for word in words:
            valor_medio = np.mean(dicionario_banco_precos[word]['valores_unitarios'])
            desvio_padrao = np.std(dicionario_banco_precos[word]['valores_unitarios'])
            valor_min_em_tese = valor_medio - desvio_padrao
            valor_minimo = valor_min_em_tese if valor_min_em_tese >= 0 else 0.00
            resultado.append(
                [
                    word,
                    'Valor médio em compras da BEC R${:.2f}'.format(valor_medio),
                    'Valor máximo e mínimo esperado em reais em compras da BEC R${:.2f} e R${:.2f}'.format(valor_medio+desvio_padrao,valor_minimo)
                ]
            )
        return resultado

    def gerar_relatorio(self):
        desc_item = self.lineEdit.text().lower()
        self.lineEdit.clear()
        resultado = self.give_recommendation(desc_item)
        documento = Document()
        for palavra, valor_medio, min_max in resultado:
            documento.add_paragraph('\n\n\nItem encontrado: '+palavra)
            documento.add_paragraph(valor_medio)
            documento.add_paragraph(min_max)
        documento.save(path_relatorio+desc_item+'.docx')
        msg = QMessageBox()
        msg.about(msg, 'Sucesso!', 'Foi gerado um relatório denominado %s.docx' % (desc_item,))


if __name__ == "__main__":
    import sys
    app = QtWidgets.QApplication(sys.argv)
    MainWindow = QtWidgets.QMainWindow()
    ui = Ui_MainWindow()
    ui.setupUi(MainWindow)
    MainWindow.show()
    sys.exit(app.exec_())
