from docx import Document
import pickle, os, re, pandas as pd, pymongo, sys

empresas_vencedoras_bl = pickle.load(open('empresas_vencedoras_bl.pickle','rb'))
ganhadores_sem_competicao = pickle.load(open('ganhadores_sem_competicao.pickle','rb'))
cnpjs_encontrados = pickle.load(open('cnpjs_encontrados_rifs.pickle','rb'))
path_rais = '/media/danilo/Seagate Expansion Drive/Dados Compras Públicas/Dicionarios_tamanho_empresas/'
arquivos_rais_n_funcionarios = [pickle.load(open(path_rais+file,'rb')) for file in os.listdir(path_rais) if file[-11:-9] == '17']
losers = pickle.load(open('losers.pickle','rb'))
licitantes_ids = pickle.load(open('licitantes_ids.pickle','rb'))
valor_contratos_licitantes = pickle.load(open('valor_contratos_licitantes.pickle','rb'))

myclient = pymongo.MongoClient("mongodb://localhost:27017/")
db = myclient["RF"]
empresas = db["Empresas"]

texto_inserir_dados_empresas = '\n\nCopie e cole na linha a seguir o número do CNPJ das empresas \nseparado por vírgulas caso sejam várias empresas e aperte enter:\n\n'

def consertar_str_cnpj(cnpj):
    cnpj = cnpj.replace('/','').replace('-','').replace('.','')
    if len(cnpj) == 14:
        return cnpj
    elif len(cnpj) > 14:
        return consertar_str_cnpj(cnpj[:-1])
    else:
        return consertar_str_cnpj('0'+cnpj)

def empresa_citada_rif(cnpj):
    if cnpj in cnpjs_encontrados:
        return 1
    return 0

def find_variaveis_receita_federal(cnpj):
    dicionario = empresas.find_one({"_id":cnpj})
    if dicionario:
        return (dicionario['situacao_cadastral'],dicionario['data_inicio_atividade'],dicionario['porte_empresa'])
    else:
        return (-1,-1,-1)

def empresa_loser(id_empresa):
    if id_empresa in losers:
        return 1
    return 0

def ganha_bl(id_empresa):
    if id_empresa in empresas_vencedoras_bl:
        return 1
    return 0

def find_n_func(cnpj):
    for arq in arquivos_rais_n_funcionarios:
        if cnpj in arq:
            return arq[cnpj]
    return -1

def sem_competicao(id_empresa):
    if int(id_empresa) in ganhadores_sem_competicao:
        return 1
    return 0

def valor_acumulado_empresa(id_empresa):
    if id_empresa in valor_contratos_licitantes:
        return valor_contratos_licitantes[id_empresa]
    else:
        return 0

def matriz_risco_empresas(lista_cnpjs,nome_relatorio='matriz_risco_empresas_licitacoes.xlsx'):
    rows = []
    for l in lista_cnpjs:
        if l in licitantes_ids:
            lic = consertar_str_cnpj(l)
            ganha = ganha_bl(licitantes_ids[lic])
            emp_loser = empresa_loser(licitantes_ids[lic])
            n_fun = find_n_func(lic)
            s_comp = sem_competicao(licitantes_ids[lic])
            e_rif = empresa_citada_rif(lic)
            v_contratos = valor_acumulado_empresa(licitantes_ids[lic])
            situacao_cadastral,data_inicio_atividade,porte_empresa = find_variaveis_receita_federal(lic)
            dicionario_row = {
                'ganha_bl' : ganha,
                'empresa_loser' : emp_loser,
                'n_funcionarios' : n_fun,
                'sem_competicao' : s_comp,
                'empresa_citada_rif' : e_rif,
                'situacao_cadastral':situacao_cadastral,
                'data_inicio_atividade':data_inicio_atividade,
                'porte_empresa':porte_empresa,
                'valor_contratos':v_contratos,
                'documento da empresa' : lic
            }
            rows.append(dicionario_row)
            gera_relatorio_empresa(ganha, emp_loser, n_fun, s_comp, v_contratos, e_rif, situacao_cadastral,data_inicio_atividade,porte_empresa,lic)
    df = pd.DataFrame(rows, index=[i for i in range(len(rows))])
    df.to_excel(nome_relatorio,index=False)

def gera_relatorio_empresa(ganha, emp_loser, n_fun, s_comp, v_contratos, e_rif, situacao_cadastral,data_inicio_atividade,porte_empresa,cnpj):
    document = Document()
    if ganha or emp_loser or (n_fun < 5) or s_comp or e_rif:
        if ganha:
            document.add_paragraph('A empresa em questão ganha de empresas consideradas "BIG LOSERS", que perdem repetidamente licitações. Isso pode ser um indicativo de que concorre contra empresas de fachada.\n\n')
            #DETALHAR PARA QUAIS EMPRESAS ELA PERDE!
        if emp_loser:
            document.add_paragraph('A empresa em questão só perdeu licitações até hoje. Isso pode ser um indicativo de seja uma empresa de fachada. Sugere-se consultar relatório da empresa para verificar de quais licitações ela participou e seus concorrentes.\n\n')
        if n_fun < 5:
            document.add_paragraph('A empresa em questão tem poucos funcionários registrados na RAIS. Ela possui %s funcionários.\n\n' % (str(n_fun),))
        if s_comp:
            document.add_paragraph('A empresa em questão ganhou as licitações sem competidores. Sugere-se consultar relatório da empresa para verificar de quais licitações ela participou.\n\n')
        if e_rif:
            document.add_paragraph('A empresa em questão aparece em um dos RIF\'s da base do Ministério Público.\n\n')
    else:
        document.add_paragraph('Não há nenhum indício de irregularidade atualmente encontrado com relação a esta empresa.')
    document.save('Relatório/relatório_geral_empresa_cnpj_%s.docx' % (cnpj,))

if __name__ == '__main__':
    matriz_risco_empresas([sys.argv[1]])